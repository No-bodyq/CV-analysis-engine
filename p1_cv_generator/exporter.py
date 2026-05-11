import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib import colors
from reportlab.lib.units import cm


# CONSTANTS
NAME_COLOR = RGBColor(28, 69, 135)       # #1C4587
HEADER_COLOR_1 = RGBColor(28, 69, 135)  # #1C4587 — top sections
HEADER_COLOR_2 = RGBColor(7, 55, 99)    # #073763 — experience/cert
EMAIL_COLOR = RGBColor(17, 85, 204)     # #1155CC
FONT = "Times New Roman"


# HELPERS

def ensure_list(value):
    if not value:
        return []
    if isinstance(value, list):
        return value
    if isinstance(value, dict):
        return [value]
    if isinstance(value, str) and value.strip():
        return [value]
    return []


def add_run(paragraph, text, size=11, bold=False, color=None, italic=False):
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(str(text))
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return run


def add_section_heading(doc, text, color=None):
    if color is None:
        color = HEADER_COLOR_1
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, text, size=11, bold=True, color=color)
    add_bottom_border(p)
    return p


def add_bottom_border(paragraph, color="1C4587"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_cell_font(cell, font_name=FONT, size=11):
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.name = font_name
            run.font.size = Pt(size)


# DOCX EXPORT
def export_docx(cv_data, user_data):
    """
    Exports CV as DOCX in exact Rukayat CV format.
    Font: Times New Roman
    Colors: Dark blue headings
    Layout: Professional Nigerian CV style
    """
    doc = Document()

    # ── Page Setup ──
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # ── Default style ──
    normal = doc.styles['Normal']
    normal.font.name = FONT
    normal.font.size = Pt(11)

    sections = cv_data.get("sections", {})
    header = sections.get("header", {})

    name = user_data.get("name", header.get("name", "")).upper()
    location = user_data.get("location", header.get("address", ""))
    email = user_data.get("email", header.get("email", ""))
    phone = user_data.get("phone", header.get("phone", ""))
    linkedin = user_data.get("linkedin", header.get("linkedin", ""))
    dob = user_data.get("dob", header.get("dob", ""))

    # NAME — Centered, Bold, Dark Blue, 24pt
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    add_run(p, name, size=24, bold=True, color=NAME_COLOR)

    # Address line
    if location:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        add_run(p, location, size=11)

    #  Email | Phone 
    if email or phone:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)

        if email:
            add_run(p, email, size=11, color=EMAIL_COLOR)
        if email and phone:
            add_run(p, " | ", size=11)
        if phone:
            add_run(p, phone, size=11)

    #  LinkedIn
    if linkedin:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        add_run(p, f"LinkedIn- {linkedin}", size=11)

    # ── DOB ──
    if dob:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        add_run(p, "DOB: ", size=11, bold=True)
        add_run(p, dob, size=11)

    # Spacer
    doc.add_paragraph()

    # PROFESSIONAL SUMMARY
    summary = sections.get("summary", "")
    if summary:
        add_section_heading(doc, "PROFESSIONAL SUMMARY")
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_run(p, summary, size=11)

    # SKILLS — Two column table with bullets
    skills = ensure_list(sections.get("skills", []))
    if skills:
        add_section_heading(doc, "SKILLS")

        table = doc.add_table(rows=0, cols=2)
        table.autofit = True

        for i in range(0, len(skills), 2):
            row_cells = table.add_row().cells

            # Left cell
            p1 = row_cells[0].paragraphs[0]
            p1.paragraph_format.space_after = Pt(2)
            add_run(p1, "• " + skills[i], size=11)
            set_cell_font(row_cells[0])

            # Right cell
            if i + 1 < len(skills):
                p2 = row_cells[1].paragraphs[0]
                p2.paragraph_format.space_after = Pt(2)
                add_run(p2, "• " + skills[i + 1], size=11)
                set_cell_font(row_cells[1])

        doc.add_paragraph()

    # EDUCATION

    education = ensure_list(sections.get("education", []))
    if education:
        add_section_heading(doc, "EDUCATION")

        for edu in education:
            if not isinstance(edu, dict):
                continue

            school = edu.get("school", "")
            degree = edu.get("degree", "")
            grade = edu.get("grade", "")
            dates = edu.get("dates", "")
            location_edu = edu.get("location", "")

            # School name — bold
            if school:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(1)

                school_text = school
                if location_edu:
                    school_text += f" — {location_edu}"
                if dates:
                    school_text += f"         {dates}"

                add_run(p, school_text, size=11, bold=True)

            # Degree
            if degree:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(1)
                add_run(p, degree, size=11)

            # Grade
            if grade:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                add_run(p, "GRADE: ", size=11, bold=True)
                add_run(p, grade, size=11)

        doc.add_paragraph()

    # WORK EXPERIENCE

    experience = ensure_list(sections.get("experience", []))
    if experience:
        add_section_heading(doc, "WORK EXPERIENCE", color=HEADER_COLOR_2)

        for job in experience:
            if not isinstance(job, dict):
                continue

            title = job.get("title", "").upper()
            company = job.get("company", "")
            location_job = job.get("location", "")
            dates = job.get("dates", "")
            bullets = job.get("bullets", [])

            # Title + Dates on same line
            if title:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(1)

                # Title bold
                r1 = p.add_run(title)
                r1.bold = True
                r1.font.name = FONT
                r1.font.size = Pt(11)

                # Dates right side with spaces
                if dates:
                    spaces = " " * max(1, 50 - len(title))
                    r2 = p.add_run(spaces + dates)
                    r2.bold = True
                    r2.font.name = FONT
                    r2.font.size = Pt(11)

            # Company + Location
            company_text = company
            if location_job and location_job not in company:
                company_text += f" — {location_job}"

            if company_text:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                add_run(p, company_text, size=11, bold=True)

            # Bullet points
            for bullet in bullets:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(18)
                p.paragraph_format.space_after = Pt(1)
                add_run(p, "• " + bullet, size=11)

            doc.add_paragraph()

 
    # CERTIFICATION & TRAINING

    certs = ensure_list(sections.get("certifications", []))
    if certs:
        add_section_heading(doc, "CERTIFICATION & TRAINING", color=HEADER_COLOR_2)

        for cert in certs:
            if not cert:
                continue

            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)

            # Format: "Organization ● Certification"
            if "●" in cert:
                parts = cert.split("●", 1)
                add_run(p, parts[0].strip(), size=11, bold=True)
                add_run(p, " ● ", size=11)
                add_run(p, parts[1].strip(), size=11)

            elif "—" in cert:
                parts = cert.split("—", 1)
                add_run(p, parts[0].strip(), size=11, bold=True)
                add_run(p, " — ", size=11)
                add_run(p, parts[1].strip(), size=11)

            else:
                add_run(p, "• " + cert, size=11, bold=True)

        doc.add_paragraph()

    # Save
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# PDF EXPORT
def export_pdf(cv_data, user_data):
    """
    Exports CV as PDF replicating Rukayat's exact format.
    """
    buffer = io.BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm
    )

    sections = cv_data.get("sections", {})
    header = sections.get("header", {})

    name = user_data.get("name", header.get("name", "")).upper()
    location = user_data.get("location", header.get("address", ""))
    email = user_data.get("email", header.get("email", ""))
    phone = user_data.get("phone", header.get("phone", ""))
    linkedin = user_data.get("linkedin", header.get("linkedin", ""))
    dob = user_data.get("dob", header.get("dob", ""))

    # Styles
    name_style = ParagraphStyle(
        'Name',
        fontName='Times-Bold',
        fontSize=20,
        alignment=TA_CENTER,
        spaceAfter=4,
        textColor=colors.HexColor('#1C4587')
    )

    contact_style = ParagraphStyle(
        'Contact',
        fontName='Times-Roman',
        fontSize=11,
        alignment=TA_CENTER,
        spaceAfter=2,
        textColor=colors.HexColor('#1a1a1a')
    )

    heading_style_1 = ParagraphStyle(
        'Heading1',
        fontName='Times-Bold',
        fontSize=11,
        alignment=TA_LEFT,
        spaceBefore=10,
        spaceAfter=2,
        textColor=colors.HexColor('#1C4587')
    )

    heading_style_2 = ParagraphStyle(
        'Heading2',
        fontName='Times-Bold',
        fontSize=11,
        alignment=TA_LEFT,
        spaceBefore=10,
        spaceAfter=2,
        textColor=colors.HexColor('#073763')
    )

    body_style = ParagraphStyle(
        'Body',
        fontName='Times-Roman',
        fontSize=11,
        alignment=TA_LEFT,
        spaceAfter=3,
        textColor=colors.HexColor('#1a1a1a'),
        leading=15
    )

    bold_style = ParagraphStyle(
        'Bold',
        fontName='Times-Bold',
        fontSize=11,
        alignment=TA_LEFT,
        spaceAfter=3,
        textColor=colors.HexColor('#1a1a1a'),
        leading=15
    )

    bullet_style = ParagraphStyle(
        'Bullet',
        fontName='Times-Roman',
        fontSize=11,
        alignment=TA_LEFT,
        spaceAfter=3,
        textColor=colors.HexColor('#1a1a1a'),
        leftIndent=15,
        leading=15
    )

    justify_style = ParagraphStyle(
        'Justify',
        fontName='Times-Roman',
        fontSize=11,
        alignment=4,  # JUSTIFY
        spaceAfter=4,
        textColor=colors.HexColor('#1a1a1a'),
        leading=15
    )

    content = []

    # Header
    content.append(Paragraph(name, name_style))
    content.append(Spacer(1, 0.2 * cm))

    if location:
        content.append(Paragraph(location, contact_style))

    contact_line = ""
    if email and phone:
        contact_line = f'<font color="#1155CC">{email}</font> | {phone}'
    elif email:
        contact_line = f'<font color="#1155CC">{email}</font>'
    elif phone:
        contact_line = phone

    if contact_line:
        content.append(Paragraph(contact_line, contact_style))

    if linkedin:
        content.append(Paragraph(f"LinkedIn- {linkedin}", contact_style))

    if dob:
        content.append(Paragraph(f"<b>DOB:</b> {dob}", contact_style))

    content.append(Spacer(1, 0.4 * cm))

    # Summary
    summary = sections.get("summary", "")
    if summary:
        content.append(Paragraph("PROFESSIONAL SUMMARY", heading_style_1))
        content.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#1C4587')))
        content.append(Spacer(1, 0.2 * cm))
        content.append(Paragraph(summary, justify_style))
        content.append(Spacer(1, 0.2 * cm))

    # Skills — two column table
    skills = sections.get("skills", [])
    if skills:
        content.append(Paragraph("SKILLS", heading_style_1))
        content.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#1C4587')))
        content.append(Spacer(1, 0.2 * cm))

        skill_rows = []
        for i in range(0, len(skills), 2):
            left = "• " + skills[i]
            right = "• " + skills[i + 1] if i + 1 < len(skills) else ""
            skill_rows.append([
                Paragraph(left, body_style),
                Paragraph(right, body_style)
            ])

        if skill_rows:
            skill_table = Table(skill_rows, colWidths=[8 * cm, 8 * cm])
            skill_table.setStyle(TableStyle([
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('LEFTPADDING', (0, 0), (-1, -1), 0),
                ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
            ]))
            content.append(skill_table)

        content.append(Spacer(1, 0.3 * cm))

    # Education
    education = sections.get("education", [])
    if education:
        content.append(Paragraph("EDUCATION", heading_style_1))
        content.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#1C4587')))
        content.append(Spacer(1, 0.2 * cm))

        for edu in education:
            if not isinstance(edu, dict):
                continue

            school = edu.get("school", "")
            degree = edu.get("degree", "")
            grade = edu.get("grade", "")
            dates = edu.get("dates", "")
            location_edu = edu.get("location", "")

            school_text = school
            if location_edu:
                school_text += f" — {location_edu}"
            if dates:
                school_text += f"    {dates}"

            if school_text:
                content.append(Paragraph(f"<b>{school_text}</b>", body_style))
            if degree:
                content.append(Paragraph(degree, body_style))
            if grade:
                content.append(Paragraph(f"<b>GRADE:</b> {grade}", body_style))

        content.append(Spacer(1, 0.3 * cm))

    # Experience
    experience = sections.get("experience", [])
    if experience:
        content.append(Paragraph("WORK EXPERIENCE", heading_style_2))
        content.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#073763')))
        content.append(Spacer(1, 0.2 * cm))

        for job in experience:
            if not isinstance(job, dict):
                continue

            title = job.get("title", "").upper()
            company = job.get("company", "")
            location_job = job.get("location", "")
            dates = job.get("dates", "")
            bullets = job.get("bullets", [])

            # Title + Dates
            if title and dates:
                title_dates_data = [[
                    Paragraph(f"<b>{title}</b>", body_style),
                    Paragraph(f"<b>{dates}</b>", body_style)
                ]]
                title_table = Table(
                    title_dates_data,
                    colWidths=[10 * cm, 6 * cm]
                )
                title_table.setStyle(TableStyle([
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 0),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                    ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
                ]))
                content.append(title_table)
            elif title:
                content.append(Paragraph(f"<b>{title}</b>", body_style))

            # Company
            company_text = company
            if location_job and location_job not in company:
                company_text += f" — {location_job}"
            if company_text:
                content.append(Paragraph(f"<b>{company_text}</b>", body_style))

            # Bullets
            for bullet in bullets:
                content.append(Paragraph("• " + bullet, bullet_style))

            content.append(Spacer(1, 0.3 * cm))

    # Certifications
    certs = sections.get("certifications", [])
    if certs:
        content.append(Paragraph("CERTIFICATION & TRAINING", heading_style_2))
        content.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#073763')))
        content.append(Spacer(1, 0.2 * cm))

        for cert in certs:
            if not cert:
                continue
            if "●" in cert:
                parts = cert.split("●", 1)
                content.append(Paragraph(
                    f"<b>{parts[0].strip()}</b> ● {parts[1].strip()}",
                    body_style
                ))
            elif "—" in cert:
                parts = cert.split("—", 1)
                content.append(Paragraph(
                    f"<b>{parts[0].strip()}</b> — {parts[1].strip()}",
                    body_style
                ))
            else:
                content.append(Paragraph(f"<b>• {cert}</b>", body_style))

    doc.build(content)
    buffer.seek(0)
    return buffer.getvalue()